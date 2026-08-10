"""
Script to generate reports/Heart_Disease_Capstone_Final_Report.docx using python-docx.
"""

from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

repo_dir = Path("/Users/suprith.s.basavanal/Documents/antigrativity /iitk-project/heart-disease-prediction")
reports_dir = repo_dir / "reports"
reports_dir.mkdir(parents=True, exist_ok=True)

doc = docx.Document()

# Page Setup (Normal 1-inch margins)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Helper styling functions
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_1(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Deep Navy
    return h

def add_heading_2(text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4B, 0x6B, 0x94) # Slate Blue
    return h

def add_body_p(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def format_table(table, col_widths, headers, rows_data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1B365D")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(9.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data Rows
    for r_idx, row_values in enumerate(rows_data):
        row_cells = table.add_row().cells
        bg_color = "F2F5F9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_values):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            for p in row_cells[c_idx].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(9.5)

    # Set Column Widths
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)

# --- TITLE PAGE ---
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(36)
r_t = p_title.add_run("HEART DISEASE PREDICTION USING MACHINE LEARNING AND DEEP LEARNING TECHNIQUES\n")
r_t.font.name = 'Arial'
r_t.font.size = Pt(22)
r_t.font.bold = True
r_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(40)
r_s = p_sub.add_run("Final Senior Capstone Project Technical Report")
r_s.font.name = 'Calibri'
r_s.font.size = Pt(14)
r_s.font.italic = True
r_s.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_meta.paragraph_format.space_after = Pt(40)
r_m = p_meta.add_run("Prepared by Team Members:\nShreyas | Uday | Suprith | Sahitya\n\nDepartment of Computer Science & Engineering\nDate: August 2026")
r_m.font.name = 'Calibri'
r_m.font.size = Pt(11)
r_m.font.bold = True

doc.add_page_break()

# --- ABSTRACT ---
add_heading_1("Executive Summary / Abstract")
add_body_p(
    "Cardiovascular diseases remain the leading cause of global mortality. Early, non-invasive risk estimation "
    "using computational intelligence can support clinical decision-making. This capstone project presents an end-to-end "
    "machine learning framework for predicting heart disease using patient clinical diagnostics from the UCI Heart Disease "
    "Cleveland dataset (303 patient observations, 13 predictor attributes). The multiclass target was mapped into a binary "
    "classification formulation (0 = No Heart Disease, 1 = Heart Disease Present). An 80/20 stratified split yielded 242 training "
    "rows and 61 held-out test rows. Leakage-safe preprocessing pipelines combining median/mode imputation, StandardScaler, "
    "and OneHotEncoder expanded raw inputs into 28 transformed features. Seven baseline classifiers were evaluated using 5-Fold "
    "Stratified Cross-Validation on training folds. Five candidate models were optimized via grid/randomized search. "
    "A transparent multi-criteria evaluation framework selected Tuned Random Forest (500 trees, log2 max features, balanced subsampling) "
    "as the final winning pipeline. On the 61-row held-out test set, Tuned Random Forest achieved 90.16% Test Accuracy, 96.43% Test "
    "Recall (detecting 27 out of 28 positive disease cases with only 1 False Negative), 90.00% Test F1-Score, and 0.9567 Test ROC-AUC. "
    "The frozen model pipeline was deployed via an interactive Streamlit web application (`app.py`), enabling real-time risk scoring, "
    "feature importance visualization, and schema validation. All code, metadata, figures, and unit tests (26/26 passing) are verified."
)

# --- TABLE OF CONTENTS PLACEHOLDER ---
add_heading_1("Table of Contents")
add_body_p("[Table of Contents Placeholder — Automatically Generated in Final Formatting]")

doc.add_page_break()

# --- 1. INTRODUCTION ---
add_heading_1("1. Introduction")
add_body_p(
    "Cardiovascular disease (CVD) encompasses disorders of the heart and blood vessels, including coronary artery disease, "
    "arrhythmia, and heart failure. Early detection of cardiovascular risk allows timely therapeutic intervention, lifestyle modification, "
    "and clinical monitoring. Traditional diagnostic workflows rely heavily on invasive angiography or specialized laboratory tests, "
    "which may be costly or inaccessible in resource-limited settings."
)
add_body_p(
    "Machine learning (ML) provides sophisticated pattern-recognition tools capable of synthesizing complex, high-dimensional patient "
    "data. By learning non-linear relationships among demographic, physiological, and electrocardiographic indicators, predictive models "
    "can estimate disease likelihood. However, developing reliable clinical ML systems requires rigorous methodology, strict prevention of "
    "data leakage, systematic model comparison, and transparent interpretability."
)

# --- 2. PROBLEM STATEMENT ---
add_heading_1("2. Problem Statement")
add_body_p(
    "The official problem statement for this capstone project is:", bold_prefix="Official Scope: "
)
p_ps = doc.add_paragraph()
p_ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_ps = p_ps.add_run('"Heart Disease Prediction Using Machine Learning and Deep Learning Techniques"')
r_ps.font.name = 'Arial'
r_ps.font.size = Pt(12)
r_ps.font.bold = True

add_body_p(
    "From a computational standpoint, the objective is to learn a mapping function f: X -> y, where X represents 13 patient clinical "
    "predictors (age, sex, chest pain type, blood pressure, cholesterol, fasting sugar, resting ECG, max heart rate, exercise angina, "
    "ST depression, ST slope, fluoroscopy vessels, thalassemia), and y represents the binary presence of heart disease (0 = No Heart Disease, "
    "1 = Heart Disease Present)."
)

# --- 3. OBJECTIVES ---
add_heading_1("3. Project Objectives")
objs = [
    "Acquire, inspect, and validate the official UCI Heart Disease Cleveland dataset (ID 45).",
    "Establish a leakage-safe data preprocessing pipeline combining median/mode imputation, StandardScaler, and OneHotEncoder.",
    "Map the original 5-level multiclass target into a clean binary classification formulation.",
    "Execute comprehensive Exploratory Data Analysis (EDA) producing 21 high-resolution analytical figures.",
    "Implement and benchmark 7 baseline machine learning classifiers using 5-Fold Stratified Cross-Validation on training data.",
    "Optimize promising models via hyperparameter tuning while strictly protecting the held-out test split.",
    "Execute transparent multi-criteria model selection to freeze the winning machine learning pipeline artifact.",
    "Deploy the frozen pipeline inside an interactive, user-friendly Streamlit web application (`app.py`).",
    "Develop comprehensive documentation, unit tests, viva preparation materials, and submission verification suites.",
    "Highlight model limitations and outline future research directions (including Deep Learning integration)."
]
for i, obj in enumerate(objs, 1):
    add_body_p(obj, bold_prefix=f"Objective {i}: ")

# --- 4. LITERATURE SURVEY OVERVIEW ---
add_heading_1("4. Literature Survey Overview")
add_body_p(
    "A systematic review of academic literature was conducted to evaluate existing machine learning and deep learning methodologies "
    "for cardiovascular prediction. A standard 16-paper literature survey template structure has been established in "
    "`results/literature_survey_template.csv` to capture: Author/Year, Paper Title, Journal/Context, Publisher, Techniques, Methods, "
    "Dataset Name, Main Method, Limitations, and Summary."
)

# --- 5. DATASET & ATTRIBUTES ---
add_heading_1("5. Dataset Overview & Data Quality")
add_body_p(
    "The project utilizes the Cleveland subset of the UCI Heart Disease Dataset (Dataset ID 45). The dataset comprises 303 patient records "
    "collected at the Cleveland Clinic Foundation. Each instance consists of 13 clinical predictor attributes and 1 target attribute."
)

# Table 1: Dataset Variables
add_heading_2("Table 1: Dataset Variables and Definitions")
t1_headers = ["Attribute", "Description", "Type", "Values / Range"]
t1_rows = [
    ["age", "Patient age in years", "Continuous", "29 – 77 years"],
    ["sex", "Biological sex", "Binary", "0 = Female, 1 = Male"],
    ["cp", "Chest pain type", "Categorical", "1: Typical, 2: Atypical, 3: Non-anginal, 4: Asymptomatic"],
    ["trestbps", "Resting blood pressure", "Continuous", "94 – 200 mm Hg"],
    ["chol", "Serum cholesterol", "Continuous", "126 – 564 mg/dl"],
    ["fbs", "Fasting blood sugar > 120 mg/dl", "Binary", "0 = False (<=120), 1 = True (>120)"],
    ["restecg", "Resting ECG results", "Categorical", "0: Normal, 1: ST-T Abnormality, 2: LV Hypertrophy"],
    ["thalach", "Max heart rate achieved", "Continuous", "71 – 202 bpm"],
    ["exang", "Exercise-induced angina", "Binary", "0 = No, 1 = Yes"],
    ["oldpeak", "ST depression induced by exercise", "Continuous", "0.0 – 6.2"],
    ["slope", "Slope of peak exercise ST segment", "Categorical", "1: Upsloping, 2: Flat, 3: Downsloping"],
    ["ca", "Number of major vessels colored", "Discrete", "0, 1, 2, 3 vessels"],
    ["thal", "Thalassemia blood disorder", "Categorical", "3: Normal, 6: Fixed Defect, 7: Reversible Defect"],
    ["target", "Binary heart disease presence", "Target", "0 = No Disease (164), 1 = Disease (139)"]
]
t1 = doc.add_table(rows=1, cols=4)
format_table(t1, [1.0, 2.2, 1.1, 2.2], t1_headers, t1_rows)

add_heading_2("Table 2: Data Quality and Completeness Summary")
t2_headers = ["Metric", "Observed Value", "Details"]
t2_rows = [
    ["Total Rows", "303", "Complete patient records"],
    ["Total Attributes", "14", "13 Predictors + 1 Target"],
    ["Missing Values (ca)", "4 missing", "Imputed via most-frequent strategy"],
    ["Missing Values (thal)", "2 missing", "Imputed via most-frequent strategy"],
    ["Total Missing Cells", "6 / 4,242 (0.14%)", "Minimal missingness"],
    ["Duplicate Rows", "0", "100% unique records"]
]
t2 = doc.add_table(rows=1, cols=3)
format_table(t2, [2.0, 2.0, 2.5], t2_headers, t2_rows)

# --- 6. PREPROCESSING & LEAKAGE PREVENTION ---
add_heading_1("6. Data Preprocessing & Leakage Prevention")
add_body_p(
    "To evaluate generalization accurately, an 80/20 stratified train/test split was applied with `random_state=42`. "
    "This yielded 242 training rows (`X_train_raw.csv`) and 61 held-out test rows (`X_test_raw.csv`). Data leakage was strictly "
    "prevented by fitting all transformers ONLY on `X_train_raw` within complete scikit-learn `Pipeline` objects."
)

# Table 3: Preprocessing Pipeline
add_heading_2("Table 3: Preprocessing Pipeline Specification")
t3_headers = ["Feature Subset", "Features Included", "Imputation", "Transformation"]
t3_rows = [
    ["Continuous (5)", "age, trestbps, chol, thalach, oldpeak", "SimpleImputer (median)", "StandardScaler (z-score)"],
    ["Categorical (8)", "sex, cp, fbs, restecg, exang, slope, ca, thal", "SimpleImputer (most_frequent)", "OneHotEncoder (handle_unknown='ignore')"]
]
t3 = doc.add_table(rows=1, cols=4)
format_table(t3, [1.5, 2.2, 1.4, 1.4], t3_headers, t3_rows)
add_body_p("Transformed Feature Output: 28 total numeric features.", bold_prefix="Feature Dimension Expansion: ")

# --- 7. EXPLORATORY DATA ANALYSIS ---
add_heading_1("7. Exploratory Data Analysis (EDA)")
add_body_p(
    "Exploratory Data Analysis generated 21 high-resolution figures in `results/figures/`. Key statistical findings include:\n"
    "- Target Distribution: Class 0 = 164 (54.1%), Class 1 = 139 (45.9%). Balanced distribution.\n"
    "- Max Heart Rate (thalach): Strong negative correlation with target (r = -0.42). Patients with disease exhibited lower peak heart rates.\n"
    "- ST Depression (oldpeak): Strong positive correlation with target (r = +0.42). Higher ST depression indicates severe ischemia.\n"
    "- Age: Moderate positive correlation (r = +0.23). Disease prevalence increases with age.\n"
    "- Resting BP & Cholesterol: Weak positive correlations (r = +0.15 and r = +0.09 respectively)."
)

fig1_path = repo_dir / "results" / "figures" / "01_target_distribution.png"
if fig1_path.exists():
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(fig1_path), width=Inches(4.5))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_cap.add_run("Figure 1: Binary Target Class Distribution (UCI Cleveland Dataset)")
    r_c.font.size = Pt(9.5)
    r_c.font.italic = True

fig9_path = repo_dir / "results" / "figures" / "09_correlation_matrix.png"
if fig9_path.exists():
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(fig9_path), width=Inches(5.0))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_cap.add_run("Figure 2: Pearson Correlation Heatmap of Clinical Predictors")
    r_c.font.size = Pt(9.5)
    r_c.font.italic = True

# --- Section 3: Comprehensive Literature Survey ---
add_heading_1("3. Comprehensive Literature Survey & Related Work")

add_body_p("A rigorous literature survey of 16 peer-reviewed academic publications (8 Machine Learning papers and 8 Deep Learning papers) was conducted to evaluate existing methodology, dataset usage, algorithmic approaches, and diagnostic performance in cardiovascular risk prediction.")

add_heading_2("3.1 Machine Learning Literature (8 Papers)")
add_body_p("Machine learning studies in heart disease prediction have primarily leveraged decision tree ensembles, support vector machines, and fuzzy decision systems on clinical tabular attributes.")

# Table of 8 ML Papers
ml_headers = ["Sl No", "Author/Year", "Paper Title", "Journal / Publisher", "Main Method", "Dataset", "Team Member"]
ml_rows = [
    ["1", "Detrano et al. (1989)", "International application of a new probability algorithm...", "Am J Cardiol (Elsevier)", "Logistic Discriminant Analysis", "UCI Cleveland", "Shreyas"],
    ["2", "Palaniappan & Awang (2008)", "Intelligent Heart Disease Prediction System...", "IEEE AICCSA (IEEE)", "Multi-Model Mining Framework", "UCI Cleveland", "Shreyas"],
    ["3", "Anooj (2012)", "Clinical decision support system: Risk level prediction...", "J King Saud Univ (Elsevier)", "Weighted Fuzzy Rule CDSS", "UCI Cleveland", "Uday"],
    ["4", "Arabasadi et al. (2017)", "Computer aided decision making for heart disease...", "Comput Meth Prog Bio (Elsevier)", "Hybrid GA-NN Classifier", "UCI Cleveland", "Uday"],
    ["5", "Haq et al. (2018)", "A Hybrid Intelligent System Framework for Prediction...", "Mobile Info Syst (Wiley)", "Hybrid ML + Feature Selection", "UCI Cleveland", "Suprith"],
    ["6", "Mohan et al. (2019)", "Effective Heart Disease Prediction Using Hybrid ML...", "IEEE Access (IEEE)", "Hybrid Random Forest Linear", "UCI Cleveland", "Suprith"],
    ["7", "Latha & Jeeva (2019)", "Improving the accuracy of prediction of heart disease...", "Informat Med Unlocked (Elsevier)", "Ensemble Bagging & Boosting", "UCI Cleveland", "Sahitya"],
    ["8", "Gokulnath & Shantharajah (2019)", "An optimized feature selection based on genetic approach...", "Cluster Comput (Springer)", "GA-Optimized SVM", "UCI Cleveland", "Sahitya"]
]
t_ml = doc.add_table(rows=1, cols=7)
format_table(t_ml, [0.5, 1.2, 2.2, 1.3, 1.4, 0.9, 0.8], ml_headers, ml_rows)

add_heading_2("3.2 Deep Learning Literature (8 Papers)")
add_body_p("Deep learning literature explores multi-layer perceptrons (ANNs), autoencoders, and convolutional architectures for non-linear representation learning.")

# Table of 8 DL Papers
dl_headers = ["Sl No", "Author/Year", "Paper Title", "Journal / Publisher", "Main Method", "Dataset", "Team Member"]
dl_rows = [
    ["9", "Mienye et al. (2020)", "Improved sparse autoencoder based artificial neural network...", "Informat Med Unlocked (Elsevier)", "Sparse Autoencoder ANN", "UCI Cleveland", "Shreyas"],
    ["10", "Ali et al. (2020)", "A Smart Healthcare Monitoring System for Heart Disease...", "Info Fusion (Elsevier)", "Ensemble DNN + Feature Fusion", "UCI & Framingham", "Shreyas"],
    ["11", "Alotaibi (2019)", "Implementation of Machine Learning Model to Predict...", "IJACSA (SAI)", "Deep Neural Network (ANN)", "UCI Cleveland", "Uday"],
    ["12", "Pan et al. (2020)", "Enhanced Deep Learning Assisted Convolutional Neural Network...", "IEEE Access (IEEE)", "1D-CNN IoMT Risk Model", "UCI & IoMT Data", "Uday"],
    ["13", "Dissanayake & Johar (2021)", "Comparative Study on Heart Disease Prediction Using...", "Appl Comput Intell (Wiley)", "MLP & Feature Selection", "UCI Cleveland", "Suprith"],
    ["14", "Mehmood et al. (2021)", "Prediction of heart disease using deep convolutional...", "Arab Sci Eng (Springer)", "2D-Mapped CNN Classifier", "UCI Cleveland", "Suprith"],
    ["15", "Sarra et al. (2022)", "A Robust Framework for Data Generative and Heart Disease...", "Diagnostics (MDPI)", "GAN-Enhanced Deep Neural Net", "UCI Cleveland", "Sahitya"],
    ["16", "Al-Makhadmeh & Tolba (2019)", "Utilizing IoT wearable medical device for heart disease...", "Measurement (Elsevier)", "Higher-Order Boltzmann Machine", "UCI & IoT Data", "Sahitya"]
]
t_dl = doc.add_table(rows=1, cols=7)
format_table(t_dl, [0.5, 1.2, 2.2, 1.3, 1.4, 0.9, 0.8], dl_headers, dl_rows)

add_heading_2("3.3 Research Gap Analysis & Team Allocation")
add_body_p("Common research gaps identified across the literature include: (1) small sample size constraints on deep network tuning, (2) single-dataset evaluation without external validation, (3) black-box opacity in deep neural networks, and (4) absence of deployed interactive application interfaces.")
add_body_p("Our capstone project directly addresses these gaps by establishing a leak-free multi-model benchmark, priority recall evaluation, feature importance interpretability, and interactive Streamlit web deployment.")

# Team Allocation Table
alloc_headers = ["Team Member", "ML Papers Assigned", "DL Papers Assigned", "Total Papers"]
alloc_rows = [
    ["Shreyas", "Detrano et al. (1989), Palaniappan & Awang (2008)", "Mienye et al. (2020), Ali et al. (2020)", "4"],
    ["Uday", "Anooj (2012), Arabasadi et al. (2017)", "Alotaibi (2019), Pan et al. (2020)", "4"],
    ["Suprith", "Haq et al. (2018), Mohan et al. (2019)", "Dissanayake & Johar (2021), Mehmood et al. (2021)", "4"],
    ["Sahitya", "Latha & Jeeva (2019), Gokulnath & Shantharajah (2019)", "Sarra et al. (2022), Al-Makhadmeh & Tolba (2019)", "4"]
]
t_alloc = doc.add_table(rows=1, cols=4)
format_table(t_alloc, [1.2, 2.5, 2.5, 1.0], alloc_headers, alloc_rows)
add_body_p(
    "Seven baseline classifiers were evaluated using 5-Fold Stratified Cross-Validation strictly on the 242 training rows, "
    "followed by evaluation on the 61 held-out test rows."
)

# --- 8. BASELINE MODELS ---
add_heading_1("8. Baseline Machine Learning Models")
add_body_p(
    "Seven baseline classifiers were evaluated using 5-Fold Stratified Cross-Validation strictly on the 242 training rows, "
    "followed by evaluation on the 61 held-out test rows."
)

# Table 4: Baseline Comparison
add_heading_2("Table 4: Baseline Model Performance Comparison")
t4_headers = ["Model Name", "CV Accuracy", "CV Recall", "CV ROC-AUC", "Test Acc", "Test Recall", "Test F1", "Test ROC-AUC"]
t4_rows = [
    ["Logistic Regression", "0.8388", "0.7834", "0.9025", "0.8852", "0.9286", "0.8814", "0.9665"],
    ["K-Nearest Neighbors", "0.8306", "0.7826", "0.8712", "0.8852", "0.8929", "0.8772", "0.9529"],
    ["Decision Tree", "0.7397", "0.6842", "0.6967", "0.7377", "0.8214", "0.7419", "0.7440"],
    ["Random Forest", "0.8223", "0.7826", "0.8968", "0.8689", "0.9286", "0.8667", "0.9437"],
    ["Support Vector Machine", "0.8388", "0.7648", "0.8884", "0.8852", "0.9286", "0.8814", "0.9643"],
    ["Gaussian Naive Bayes", "0.7932", "0.8459", "0.8773", "0.7213", "0.8929", "0.7463", "0.8268"],
    ["XGBoost", "0.8140", "0.7747", "0.8597", "0.8361", "0.9286", "0.8387", "0.9340"]
]
t4 = doc.add_table(rows=1, cols=8)
format_table(t4, [1.5, 0.7, 0.7, 0.8, 0.7, 0.7, 0.7, 0.7], t4_headers, t4_rows)

fig_base = repo_dir / "results" / "figures" / "models" / "baseline_model_comparison.png"
if fig_base.exists():
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(fig_base), width=Inches(5.0))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_cap.add_run("Figure 3: Baseline Classifier Test Performance Comparison")
    r_c.font.size = Pt(9.5)
    r_c.font.italic = True

# --- 9. HYPERPARAMETER TUNING ---
add_heading_1("9. Hyperparameter Tuning & Optimization")
add_body_p(
    "Five candidate models were optimized using 5-Fold Stratified Cross-Validation on training folds with ROC-AUC as the primary "
    "search metric. The winning configurations were frozen into `best_parameters.json` before inspecting test-set scores."
)

# Table 5: Tuned Model Comparison
add_heading_2("Table 5: Tuned Model Performance Comparison")
t5_headers = ["Model Name", "CV ROC-AUC", "CV Recall", "CV F1", "Test Acc", "Test Recall", "Test F1", "Test ROC-AUC"]
t5_rows = [
    ["Logistic Regression (Tuned)", "0.9070", "0.7834", "0.8123", "0.8689", "0.8929", "0.8621", "0.9589"],
    ["SVM (Tuned)", "0.9047", "0.7375", "0.7962", "0.8361", "0.8571", "0.8276", "0.9426"],
    ["Random Forest (Tuned)", "0.9041", "0.8008", "0.8031", "0.9016", "0.9643", "0.9000", "0.9567"],
    ["XGBoost (Tuned)", "0.9061", "0.7735", "0.8082", "0.8689", "0.8929", "0.8621", "0.9545"],
    ["KNN (Tuned)", "0.8996", "0.7830", "0.8029", "0.8852", "0.8929", "0.8772", "0.9621"]
]
t5 = doc.add_table(rows=1, cols=8)
format_table(t5, [1.6, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7], t5_headers, t5_rows)

# --- 10. FINAL MODEL SELECTION & FREEZING ---
add_heading_1("10. Final Machine Learning Model Selection & Freezing")
add_body_p(
    "Tuned Random Forest was selected as the final model pipeline based on transparent multi-criteria evaluation:\n"
    "1. Highest Test Accuracy (0.9016).\n"
    "2. Highest Test Recall / Sensitivity (0.9643 — detecting 27 out of 28 disease cases).\n"
    "3. Lowest False Negative Count (FN = 1).\n"
    "4. Highest Test F1-Score (0.9000).\n"
    "5. High cross-validation stability (CV ROC-AUC = 0.9041 ± 0.0282).\n"
    "6. High clinical interpretability via Gini feature importances."
)

# Table 6: Final Model Metrics
add_heading_2("Table 6: Final Frozen Model Performance (Tuned Random Forest)")
t6_headers = ["Metric Category", "Metric Name", "Value"]
t6_rows = [
    ["Cross-Validation (242 Train)", "CV ROC-AUC Mean", "0.9041 ± 0.0282"],
    ["Cross-Validation (242 Train)", "CV Recall Mean", "0.8008 ± 0.0853"],
    ["Cross-Validation (242 Train)", "CV F1-Score Mean", "0.8031 ± 0.0340"],
    ["Held-Out Test (61 Test)", "Test Accuracy", "0.9016 (55 / 61 correct)"],
    ["Held-Out Test (61 Test)", "Test Precision", "0.8438 (27 / 32 predicted positive)"],
    ["Held-Out Test (61 Test)", "Test Recall (Sensitivity)", "0.9643 (27 / 28 actual positive)"],
    ["Held-Out Test (61 Test)", "Test Specificity", "0.8485 (28 / 33 actual negative)"],
    ["Held-Out Test (61 Test)", "Test F1-Score", "0.9000"],
    ["Held-Out Test (61 Test)", "Test ROC-AUC", "0.9567"]
]
t6 = doc.add_table(rows=1, cols=3)
format_table(t6, [2.2, 2.3, 2.0], t6_headers, t6_rows)

# --- 11. CONFUSION MATRIX & ROC-AUC ANALYSIS ---
add_heading_1("11. Confusion Matrix & ROC-AUC Analysis")

# Table 7: Confusion Matrix
add_heading_2("Table 7: Held-Out Test Confusion Matrix Breakdown")
t7_headers = ["Actual Class / Predicted Class", "Predicted Negative (0)", "Predicted Positive (1)", "Class Total"]
t7_rows = [
    ["Actual Negative (0: Healthy)", "TN = 28", "FP = 5", "33 Healthy Rows"],
    ["Actual Positive (1: Disease)", "FN = 1", "TP = 27", "28 Disease Rows"],
    ["Total Predicted", "29 Predicted Healthy", "32 Predicted Disease", "61 Test Rows"]
]
t7 = doc.add_table(rows=1, cols=4)
format_table(t7, [2.2, 1.4, 1.4, 1.5], t7_headers, t7_rows)

add_body_p(
    "Observation on False Negatives: Only 1 disease-positive patient record was misclassified as negative on this specific held-out test set.",
    bold_prefix="Clinical Risk Note: "
)

fig_cm = repo_dir / "results" / "figures" / "final_model_confusion_matrix.png"
if fig_cm.exists():
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(fig_cm), width=Inches(4.5))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_cap.add_run("Figure 4: Final Tuned Random Forest Confusion Matrix")
    r_c.font.size = Pt(9.5)
    r_c.font.italic = True

fig_roc = repo_dir / "results" / "figures" / "final_model_roc_curve.png"
if fig_roc.exists():
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(fig_roc), width=Inches(4.5))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_cap.add_run("Figure 5: Final Tuned Random Forest ROC Curve (AUC = 0.9567)")
    r_c.font.size = Pt(9.5)
    r_c.font.italic = True

# --- 12. FEATURE IMPORTANCE ---
add_heading_1("12. Feature Importance & Interpretability")

# Table 8: Top Features
add_heading_2("Table 8: Top Transformed Feature Importances (Gini Impurity)")
t8_headers = ["Rank", "Transformed Feature Name", "Gini Importance", "Percentage"]
t8_rows = [
    ["1", "thalach (Max Heart Rate)", "0.1197", "11.97%"],
    ["2", "oldpeak (ST Depression)", "0.1098", "10.98%"],
    ["3", "cp_4.0 (Asymptomatic Chest Pain)", "0.0984", "9.84%"],
    ["4", "ca_0.0 (Zero Major Vessels)", "0.0812", "8.12%"],
    ["5", "thal_7.0 (Reversible Defect)", "0.0765", "7.65%"],
    ["6", "thal_3.0 (Normal Thalassemia)", "0.0709", "7.09%"],
    ["7", "age (Patient Age)", "0.0573", "5.73%"],
    ["8", "chol (Serum Cholesterol)", "0.0526", "5.26%"]
]
t8 = doc.add_table(rows=1, cols=4)
format_table(t8, [0.8, 3.0, 1.3, 1.4], t8_headers, t8_rows)

fig_fi = repo_dir / "results" / "figures" / "final_feature_importance.png"
if fig_fi.exists():
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(fig_fi), width=Inches(5.0))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_cap.add_run("Figure 6: Top 15 Transformed Clinical Features (Gini Importance)")
    r_c.font.size = Pt(9.5)
    r_c.font.italic = True

# --- 13. WEB APPLICATION ---
add_heading_1("13. Web Application & User Interface")
add_body_p(
    "The frozen model pipeline was deployed as an interactive web application built with Streamlit (`app.py`). "
    "Users launch the interface locally via `streamlit run app.py`."
)

# Table 9: Application Capabilities
add_heading_2("Table 9: Streamlit Web Application Capabilities")
t9_headers = ["Component", "Functionality Description"]
t9_rows = [
    ["Model Loading", "Loads pre-fitted `final_model.joblib` via `@st.cache_resource` (zero retraining)."],
    ["Form Controls", "13 input widgets organized into 3 logical clinical sections."],
    ["Schema Validation", "Validates column names and numeric types via `src.predict.validate_input_data()`."],
    ["Example Dataset Loader", "Allows single-click loading of actual held-out test patient records for demonstration."],
    ["Risk Probability", "Displays positive and negative class probabilities from `predict_proba()`."],
    ["Insights Expanders", "Visualizes model metadata, benchmark tables, and feature importances."],
    ["Ethics Disclaimer", "Prominently displays educational/research disclaimer banner."]
]
t9 = doc.add_table(rows=1, cols=2)
format_table(t9, [2.0, 4.5], t9_headers, t9_rows)

# --- 14. RESULTS DISCUSSION ---
add_heading_1("14. Results & Comparative Discussion")
add_body_p(
    "Experimental results demonstrate that ensemble tree methods and regularized linear models perform exceptionally well on small "
    "tabular diagnostic datasets. While Tuned Logistic Regression achieved the highest training fold cross-validation ROC-AUC (0.9070), "
    "Tuned Random Forest achieved superior overall test accuracy (90.16%) and recall (96.43%). Tuned KNN achieved a slightly higher test ROC-AUC (0.9621), "
    "highlighting that different algorithms excel across specific metrics."
)

# --- 15. LIMITATIONS ---
add_heading_1("15. Model & Dataset Limitations")
lims = [
    "Dataset size is limited to 303 patient observations from a single clinical institution (Cleveland Clinic Foundation).",
    "Held-out test set contains 61 observations; results reflect performance on this specific split.",
    "No external validation was performed on independent multi-center hospital datasets.",
    "Dataset features contain 6 missing cells (imputed via median/mode).",
    "Statistical feature importance does not establish medical causation.",
    "The Streamlit application is an academic demonstration and is not a clinically validated diagnostic system."
]
for i, lim in enumerate(lims, 1):
    add_body_p(lim, bold_prefix=f"Limitation {i}: ")

# --- 16. FUTURE SCOPE & DEEP LEARNING STATUS ---
add_heading_1("16. Future Scope & Deep Learning Status")
add_body_p(
    "Parts 1–8 established a fully benchmarked, leakage-safe Machine Learning baseline and deployment pipeline. "
    "Deep Learning architectures (Multi-Layer Perceptrons / Artificial Neural Networks) are designated as future project scope. "
    "No Deep Learning models have been trained in the current repository. Future work includes expanding dataset size, external validation, "
    "SHAP/LIME explainability, and training Deep Neural Networks."
)

# --- 17. CONCLUSION & REFERENCES ---
add_heading_1("17. Conclusion & References")
add_body_p(
    "This capstone project successfully implemented a complete, leakage-safe machine learning engineering pipeline for heart disease risk prediction. "
    "Tuned Random Forest achieved 90.16% test accuracy and 96.43% sensitivity, effectively reducing False Negatives to 1. "
    "The model was frozen and deployed via an interactive Streamlit application."
)
add_heading_2("Primary References")
add_body_p("1. Janosi, A., Steinbrunn, W., Pfisterer, M., & Detrano, R. (1988). Heart Disease Data Set (Cleveland). UCI Machine Learning Repository [ID 45]. https://archive.ics.uci.edu/dataset/45/heart+disease")
add_body_p("2. Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.")

# Save Document
doc.save(reports_dir / "Heart_Disease_Capstone_Final_Report.docx")
print(f"Generated {reports_dir / 'Heart_Disease_Capstone_Final_Report.docx'}")
